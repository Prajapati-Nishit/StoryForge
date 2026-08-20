import io

from django.http import HttpResponse


def _story_sections(story):
    sections = [
        ("Title", story.title),
        ("Genre / Theme", f"{story.genre} / {story.theme}"),
        ("Tagline", story.tagline),
        ("Summary", story.summary),
    ]
    return sections


def _safe_filename(story):
    return "".join(c if c.isalnum() or c in " -_" else "" for c in story.title).strip().replace(" ", "_") or "story"


def build_markdown(story):
    lines = [f"# {story.title}", ""]
    lines.append(f"*{story.tagline}*")
    lines.append("")
    lines.append(f"**Genre:** {story.genre}  **Theme:** {story.theme}  **Tone:** {story.tone}")
    lines.append("")
    lines.append("## Summary")
    lines.append(story.summary)
    lines.append("")
    lines.append("## Characters")
    for c in story.characters.all():
        lines.append(f"### {c.name} ({c.age}) — {c.role}")
        lines.append(f"- Occupation: {c.occupation}")
        lines.append(f"- Personality: {c.personality}")
        lines.append(f"- Goal: {c.goal}")
        lines.append(f"- Weakness: {c.weakness}")
        lines.append(f"- Strength: {c.strength}")
        lines.append(f"- Skills: {c.skills}")
        lines.append(f"- Backstory: {c.backstory}")
        lines.append("")
    world = getattr(story, "world", None)
    if world:
        lines.append("## World")
        lines.append(f"**{world.name}** — {world.description}")
        lines.append(f"- Kingdoms / cities: {world.kingdoms_cities}")
        lines.append(f"- Magic / technology: {world.magic_or_technology}")
        lines.append(f"- Climate: {world.climate}")
        lines.append(f"- Culture: {world.culture}")
        lines.append(f"- History: {world.history}")
        lines.append("")
    lines.append("## Story structure")
    lines.append(f"- **Beginning:** {story.structure_beginning}")
    lines.append(f"- **Conflict:** {story.structure_conflict}")
    lines.append(f"- **Rising action:** {story.structure_rising_action}")
    lines.append(f"- **Climax:** {story.structure_climax}")
    lines.append(f"- **Ending:** {story.structure_ending}")
    lines.append("")
    lines.append("## Chapters")
    for ch in story.chapters.all():
        lines.append(f"### Chapter {ch.order}: {ch.title}")
        lines.append(ch.summary)
        lines.append("")
    lines.append("## Dialogue")
    for d in story.dialogue_lines.all():
        lines.append(f"**{d.character_name}:** {d.line}")
    lines.append("")
    lines.append("## Plot twist")
    lines.append(story.plot_twist)
    lines.append("")
    lines.append("## Ending")
    lines.append(story.ending)
    lines.append("")
    lines.append("## Sequel ideas")
    for s in story.sequel_ideas.all():
        lines.append(f"- {s.text}")
    return "\n".join(lines)


def export_markdown(story):
    content = build_markdown(story)
    resp = HttpResponse(content, content_type="text/markdown")
    resp["Content-Disposition"] = f'attachment; filename="{_safe_filename(story)}.md"'
    return resp


def export_txt(story):
    content = build_markdown(story)
    for token in ["# ", "## ", "### ", "**", "*"]:
        content = content.replace(token, "")
    resp = HttpResponse(content, content_type="text/plain")
    resp["Content-Disposition"] = f'attachment; filename="{_safe_filename(story)}.txt"'
    return resp


def export_docx(story):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(story.title, level=0)
    if story.tagline:
        p = doc.add_paragraph()
        run = p.add_run(story.tagline)
        run.italic = True
    doc.add_paragraph(f"Genre: {story.genre}    Theme: {story.theme}    Tone: {story.tone}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(story.summary)

    doc.add_heading("Characters", level=1)
    for c in story.characters.all():
        doc.add_heading(f"{c.name} ({c.age}) — {c.role}", level=2)
        for label, val in [
            ("Occupation", c.occupation), ("Personality", c.personality),
            ("Goal", c.goal), ("Weakness", c.weakness), ("Strength", c.strength),
            ("Skills", c.skills), ("Backstory", c.backstory),
        ]:
            doc.add_paragraph(f"{label}: {val}")

    world = getattr(story, "world", None)
    if world:
        doc.add_heading("World", level=1)
        doc.add_paragraph(f"{world.name} — {world.description}")
        for label, val in [
            ("Kingdoms / cities", world.kingdoms_cities),
            ("Magic / technology", world.magic_or_technology),
            ("Climate", world.climate), ("Culture", world.culture),
            ("History", world.history),
        ]:
            doc.add_paragraph(f"{label}: {val}")

    doc.add_heading("Story structure", level=1)
    for label, val in [
        ("Beginning", story.structure_beginning), ("Conflict", story.structure_conflict),
        ("Rising action", story.structure_rising_action), ("Climax", story.structure_climax),
        ("Ending", story.structure_ending),
    ]:
        doc.add_paragraph(f"{label}: {val}")

    doc.add_heading("Chapters", level=1)
    for ch in story.chapters.all():
        doc.add_heading(f"Chapter {ch.order}: {ch.title}", level=2)
        doc.add_paragraph(ch.summary)

    doc.add_heading("Dialogue", level=1)
    for d in story.dialogue_lines.all():
        doc.add_paragraph(f"{d.character_name}: {d.line}")

    doc.add_heading("Plot twist", level=1)
    doc.add_paragraph(story.plot_twist)

    doc.add_heading("Ending", level=1)
    doc.add_paragraph(story.ending)

    doc.add_heading("Sequel ideas", level=1)
    for s in story.sequel_ideas.all():
        doc.add_paragraph(f"- {s.text}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{_safe_filename(story)}.docx"'
    return resp


def export_pdf(story):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("StoryTitle", parent=styles["Title"], spaceAfter=6)
    h_style = ParagraphStyle("H", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
    body = styles["BodyText"]

    story_flow = [Paragraph(story.title, title_style)]
    if story.tagline:
        story_flow.append(Paragraph(f"<i>{story.tagline}</i>", body))
    story_flow.append(Paragraph(f"Genre: {story.genre} &nbsp;&nbsp; Theme: {story.theme} &nbsp;&nbsp; Tone: {story.tone}", body))

    story_flow.append(Paragraph("Summary", h_style))
    story_flow.append(Paragraph(story.summary, body))

    story_flow.append(Paragraph("Characters", h_style))
    for c in story.characters.all():
        story_flow.append(Paragraph(f"<b>{c.name}</b> ({c.age}) — {c.role}", body))
        story_flow.append(Paragraph(
            f"Occupation: {c.occupation} | Personality: {c.personality} | Goal: {c.goal} | "
            f"Weakness: {c.weakness} | Strength: {c.strength} | Skills: {c.skills}", body))
        story_flow.append(Spacer(1, 6))

    world = getattr(story, "world", None)
    if world:
        story_flow.append(Paragraph("World", h_style))
        story_flow.append(Paragraph(f"<b>{world.name}</b> — {world.description}", body))
        story_flow.append(Paragraph(
            f"Kingdoms/cities: {world.kingdoms_cities} | Magic/technology: {world.magic_or_technology} | "
            f"Climate: {world.climate} | Culture: {world.culture}", body))

    story_flow.append(Paragraph("Story structure", h_style))
    for label, val in [
        ("Beginning", story.structure_beginning), ("Conflict", story.structure_conflict),
        ("Rising action", story.structure_rising_action), ("Climax", story.structure_climax),
        ("Ending", story.structure_ending),
    ]:
        story_flow.append(Paragraph(f"<b>{label}:</b> {val}", body))

    story_flow.append(Paragraph("Chapters", h_style))
    for ch in story.chapters.all():
        story_flow.append(Paragraph(f"<b>Chapter {ch.order}: {ch.title}</b>", body))
        story_flow.append(Paragraph(ch.summary, body))
        story_flow.append(Spacer(1, 4))

    story_flow.append(Paragraph("Dialogue", h_style))
    for d in story.dialogue_lines.all():
        story_flow.append(Paragraph(f"<b>{d.character_name}:</b> {d.line}", body))

    story_flow.append(Paragraph("Plot twist", h_style))
    story_flow.append(Paragraph(story.plot_twist, body))

    story_flow.append(Paragraph("Ending", h_style))
    story_flow.append(Paragraph(story.ending, body))

    story_flow.append(Paragraph("Sequel ideas", h_style))
    for s in story.sequel_ideas.all():
        story_flow.append(Paragraph(f"- {s.text}", body))

    doc.build(story_flow)
    buf.seek(0)
    resp = HttpResponse(buf.read(), content_type="application/pdf")
    resp["Content-Disposition"] = f'attachment; filename="{_safe_filename(story)}.pdf"'
    return resp
